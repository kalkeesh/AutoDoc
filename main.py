import os
import math
import shutil
import queue
import json
import tkinter as tk
import threading
import ctypes
from ctypes import wintypes
from datetime import datetime
from tkinter import filedialog, messagebox

import pyautogui
from PIL import Image, ImageDraw, ImageTk
from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.text.paragraph import Paragraph

APP_BG = "#f3f3f3"
PRIMARY_BLUE = "#4f6ff0"
PRIMARY_RED = "#f61d2a"
BUTTON_SHADOW = "#d3d3d3"
SCREENSHOT_DIR = "screenshots"
PREVIEW_NOTES_FILE = "preview_notes.json"
CARD_BG = "#ffffff"
TEXT_DARK = "#20304a"
BASE_WIDTH = 420
BASE_HEIGHT = 680
MIN_SCALE = 0.70
MAX_SCALE = 1.25
TRANSPARENT_KEY = "#00ff00"
SHOT_SHORTCUT_LABEL = "Ctrl+Alt+Q"
WORKAREA_SHORTCUT_LABEL = "Ctrl+Alt+W"
PREVIEW_BRUSH_LEVELS = [8, 14, 22, 32]
SHOT_SINGLE_CLICK_DELAY_MS = 420
HOTKEY_ID = 1
WORKAREA_HOTKEY_ID = 2
MOD_ALT = 0x0001
MOD_CONTROL = 0x0002
SPI_GETWORKAREA = 0x0030
WM_HOTKEY = 0x0312
WM_QUIT = 0x0012
VK_Q = 0x51
VK_W = 0x57


def ensure_screenshot_dir():
    os.makedirs(SCREENSHOT_DIR, exist_ok=True)
    return SCREENSHOT_DIR


def ensure_screenshot_backup_dir():
    backup_dir = os.path.join(ensure_screenshot_dir(), ".originals")
    os.makedirs(backup_dir, exist_ok=True)
    return backup_dir


def get_screenshot_backup_path(image_path):
    return os.path.join(ensure_screenshot_backup_dir(), os.path.basename(image_path))


def ensure_screenshot_backup(image_path):
    backup_path = get_screenshot_backup_path(image_path)
    if os.path.exists(image_path) and not os.path.exists(backup_path):
        shutil.copyfile(image_path, backup_path)
    return backup_path


def get_preview_notes_path():
    return os.path.join(ensure_screenshot_dir(), PREVIEW_NOTES_FILE)


def load_preview_notes():
    notes_path = get_preview_notes_path()
    if not os.path.exists(notes_path):
        return {}
    try:
        with open(notes_path, "r", encoding="utf-8") as file:
            data = json.load(file)
    except (OSError, json.JSONDecodeError):
        return {}
    return data if isinstance(data, dict) else {}


def save_preview_notes(notes):
    notes_path = get_preview_notes_path()
    try:
        with open(notes_path, "w", encoding="utf-8") as file:
            json.dump(notes, file, ensure_ascii=False, indent=2)
    except OSError:
        pass


def set_preview_note(image_path, note_text):
    if not image_path:
        return
    notes = load_preview_notes()
    key = os.path.basename(image_path)
    if note_text:
        notes[key] = note_text
    else:
        notes.pop(key, None)
    save_preview_notes(notes)


def get_preview_note(image_path):
    if not image_path:
        return ""
    return load_preview_notes().get(os.path.basename(image_path), "")


def remove_preview_note(image_path):
    if not image_path:
        return
    notes = load_preview_notes()
    if notes.pop(os.path.basename(image_path), None) is not None:
        save_preview_notes(notes)


def clear_screenshot_dir():
    screenshot_dir = ensure_screenshot_dir()
    removed_count = 0

    for name in os.listdir(screenshot_dir):
        path = os.path.join(screenshot_dir, name)
        if not os.path.isfile(path):
            continue
        if name == PREVIEW_NOTES_FILE:
            os.remove(path)
            continue
        if not name.lower().endswith((".png", ".jpg", ".jpeg")):
            continue
        os.remove(path)
        removed_count += 1

    backup_dir = ensure_screenshot_backup_dir()
    for name in os.listdir(backup_dir):
        path = os.path.join(backup_dir, name)
        if not os.path.isfile(path):
            continue
        if not name.lower().endswith((".png", ".jpg", ".jpeg")):
            continue
        os.remove(path)

    return removed_count


def list_saved_screenshots():
    screenshot_dir = ensure_screenshot_dir()
    paths = []
    for name in os.listdir(screenshot_dir):
        lower_name = name.lower()
        if lower_name.endswith((".png", ".jpg", ".jpeg")):
            paths.append(os.path.join(screenshot_dir, name))
    return sorted(paths, key=os.path.getmtime)


def get_workarea_bbox():
    if os.name != "nt":
        screenshot = pyautogui.screenshot()
        return 0, 0, screenshot.width, screenshot.height

    rect = wintypes.RECT()
    ctypes.windll.user32.SystemParametersInfoW(SPI_GETWORKAREA, 0, ctypes.byref(rect), 0)
    return rect.left, rect.top, rect.right, rect.bottom


def capture_image(exclude_taskbar=False):
    screenshot = pyautogui.screenshot()
    if not exclude_taskbar:
        return screenshot

    left, top, right, bottom = get_workarea_bbox()
    return screenshot.crop((left, top, right, bottom))


def clean_filename(value):
    cleaned = "".join("_" if char in '<>:"/\\|?*' else char for char in value).strip(" .")
    return cleaned or datetime.now().strftime("TestDoc_%Y%m%d_%H%M%S")


class DocumentManager:
    def __init__(self, app):
        self.app = app
        self.document = None
        self.path = None
        self.tracked_images = {}
        self.tracked_notes = {}

    def has_document(self):
        return self.document is not None and self.path is not None

    def save(self):
        if not self.has_document():
            return
        try:
            self.document.save(self.path)
        except Exception as exc:
            messagebox.showerror("Save failed", f"Could not save the document.\n\n{exc}")

    def select_document(self):
        file_path = filedialog.askopenfilename(filetypes=[("Word Files", "*.docx")])
        if not file_path:
            return

        try:
            self.document = Document(file_path)
            self.path = file_path
            self.tracked_images.clear()
            self.tracked_notes.clear()
            imported_count = self.import_document_preview_items()
            if imported_count:
                self.app.preview_manager.refresh_paths()
                self.app.set_status(f"Selected: {os.path.basename(file_path)}. Loaded {imported_count} preview image(s).")
            else:
                self.app.set_status(f"Selected: {os.path.basename(file_path)}")
        except Exception as exc:
            messagebox.showerror("Open failed", f"Could not open document.\n\n{exc}")

    def create_new_document_file(self, folder, title):
        filename = os.path.join(folder, f"{clean_filename(title)}.docx")
        try:
            self.document = Document()
            self.document.add_heading(title, 0)
            self.document.save(filename)
            self.path = os.path.abspath(filename)
            self.tracked_images.clear()
            self.tracked_notes.clear()
            clear_screenshot_dir()
            self.app.preview_manager.refresh_paths()
            self.app.set_status(f"Created: {os.path.basename(self.path)}")
            return True
        except Exception as exc:
            messagebox.showerror("Create failed", f"Could not create document.\n\n{exc}")
            return False

    def _normalize_path(self, image_path):
        return os.path.abspath(image_path) if image_path else None

    def append_image(self, image_path, width=Inches(5), note_element=None):
        if not self.has_document() or not image_path:
            return

        absolute_path = self._normalize_path(image_path)
        if absolute_path in self.tracked_images:
            self.remove_tracked_image(absolute_path)

        picture = self.document.add_picture(image_path, width=width)
        paragraph = self.document.paragraphs[-1]
        self.tracked_images[absolute_path] = paragraph._element
        if note_element is not None:
            self.tracked_notes[absolute_path] = note_element

    def remove_tracked_image(self, image_path, remove_note=True):
        absolute_path = self._normalize_path(image_path)
        if absolute_path is None:
            return False

        element = self.tracked_images.get(absolute_path)
        if element is None:
            return False

        parent = element.getparent()
        if parent is None:
            self.tracked_images.pop(absolute_path, None)
            return False

        parent.remove(element)
        self.tracked_images.pop(absolute_path, None)
        if remove_note:
            note_element = self.tracked_notes.pop(absolute_path, None)
            if note_element is not None:
                note_parent = note_element.getparent()
                if note_parent is not None:
                    note_parent.remove(note_element)
        return True

    def add_note_runs(self, runs):
        if not self.has_document() or not runs:
            return None

        paragraph = self.document.add_paragraph()
        for text, style in runs:
            if not text:
                continue
            run = paragraph.add_run(text)
            if style == "bold":
                run.bold = True
            elif style == "italic":
                run.italic = True
            elif style == "bold_italic":
                run.bold = True
                run.italic = True
        return paragraph._element

    def set_image_note(self, image_path, note_text):
        if not self.has_document() or not image_path:
            return

        absolute_path = self._normalize_path(image_path)
        note_element = self.tracked_notes.get(absolute_path)
        if note_element is None and note_text:
            image_element = self.tracked_images.get(absolute_path)
            if image_element is not None:
                note_element = OxmlElement("w:p")
                image_element.addprevious(note_element)
                self.tracked_notes[absolute_path] = note_element

        if note_element is None:
            return

        paragraph = Paragraph(note_element, self.document)
        paragraph.clear()
        if note_text:
            paragraph.add_run(note_text)
        self.save()

    def import_document_preview_items(self):
        if not self.has_document():
            return 0

        clear_screenshot_dir()
        imported_count = 0
        pending_note_parts = []

        for paragraph in self.document.paragraphs:
            image_rids = paragraph._element.xpath(".//a:blip/@r:embed")
            if not image_rids:
                text = paragraph.text.strip()
                if text:
                    pending_note_parts.append(text)
                continue

            note_text = "\n".join(pending_note_parts).strip()
            pending_note_parts = []

            for rid in image_rids:
                related_part = self.document.part.related_parts.get(rid)
                if related_part is None:
                    continue
                extension = os.path.splitext(related_part.partname)[1] or ".png"
                imported_count += 1
                image_name = f"doc_image_{imported_count:03d}{extension}"
                image_path = os.path.join(ensure_screenshot_dir(), image_name)
                try:
                    with open(image_path, "wb") as image_file:
                        image_file.write(related_part.blob)
                    ensure_screenshot_backup(image_path)
                except OSError:
                    continue

                absolute_path = self._normalize_path(image_path)
                self.tracked_images[absolute_path] = paragraph._element
                if note_text:
                    self.tracked_notes[absolute_path] = self._find_note_element_before(paragraph._element)
                    set_preview_note(image_path, note_text)

        return imported_count

    def _find_note_element_before(self, image_element):
        previous = image_element.getprevious()
        while previous is not None:
            paragraph = Paragraph(previous, self.document)
            if paragraph.text.strip():
                return previous
            previous = previous.getprevious()
        return None


class NoteManager:
    def __init__(self, app, document_manager):
        self.app = app
        self.document_manager = document_manager
        self.current_note_text = ""
        self.window = None
        self.textbox = None
        self.bold_button = None
        self.italic_button = None

    def get_note_text(self):
        if self.window and self.window.winfo_exists() and self.textbox:
            return self.textbox.get("1.0", tk.END).strip()
        return self.current_note_text

    def save(self, close_after=False):
        self.current_note_text = self.get_note_text()
        if self.current_note_text:
            self.app.set_status("Note saved. Press Shot to use it.")
        else:
            self.app.set_status("Note cleared.")

        if close_after:
            self.hide()

    def clear(self):
        self.current_note_text = ""
        if self.textbox is not None:
            self.textbox.delete("1.0", tk.END)
            for tag_name in ("bold", "italic", "bold_italic"):
                self.textbox.tag_remove(tag_name, "1.0", tk.END)
        self.update_style_buttons()
        self.app.set_status("Note cleared.")

    def is_visible(self):
        return self.window is not None and self.window.winfo_exists() and self.window.state() != "withdrawn"

    def hide(self):
        if self.window is not None and self.window.winfo_exists():
            self.window.withdraw()

    def show(self):
        if self.window is not None and self.window.winfo_exists():
            self.window.deiconify()
            self.window.lift()
            self.window.focus_force()

    def open_window(self):
        if self.window is not None and self.window.winfo_exists():
            self.show()
            return

        self.window = tk.Toplevel(self.app.root)
        self.window.title("Add Note")
        self.window.geometry("420x320+1020+170")
        self.window.configure(bg=CARD_BG)
        self.window.attributes("-topmost", True)

        header = tk.Label(
            self.window,
            text="Write your note",
            bg=CARD_BG,
            fg=TEXT_DARK,
            font=("Segoe UI Semibold", 12),
        )
        header.pack(anchor="w", padx=16, pady=(14, 6))

        toolbar = tk.Frame(self.window, bg=CARD_BG)
        toolbar.pack(fill="x", padx=16, pady=(0, 8))

        self.bold_button = tk.Button(toolbar, text="Bold", width=8, command=self.toggle_bold, relief="flat")
        self.bold_button.pack(side="left", padx=(0, 8))

        self.italic_button = tk.Button(toolbar, text="Italic", width=8, command=self.toggle_italic, relief="flat")
        self.italic_button.pack(side="left", padx=(0, 8))

        clear_button = tk.Button(
            toolbar,
            text="Clear",
            width=8,
            command=self.clear,
            relief="flat",
            bg="#ffe4e6",
            fg="#8b1e2d",
        )
        clear_button.pack(side="left")

        self.textbox = tk.Text(
            self.window,
            wrap="word",
            font=("Segoe UI", 10),
            bd=1,
            relief="solid",
            padx=10,
            pady=10,
            fg=TEXT_DARK,
            highlightthickness=1,
            highlightbackground="#d9deea",
        )
        self.textbox.pack(fill="both", expand=True, padx=16, pady=(0, 12))
        self.textbox.insert("1.0", self.current_note_text)
        self.configure_text_tags()
        self.textbox.bind("<<Selection>>", lambda _event: self.update_style_buttons())
        self.textbox.bind("<ButtonRelease-1>", lambda _event: self.update_style_buttons())
        self.textbox.bind("<KeyRelease>", lambda _event: self.update_style_buttons())

        action_bar = tk.Frame(self.window, bg=CARD_BG)
        action_bar.pack(fill="x", padx=16, pady=(0, 14))

        save_button = tk.Button(
            action_bar,
            text="Save & Close",
            command=lambda: self.save(close_after=True),
            bg=PRIMARY_BLUE,
            fg="white",
            relief="flat",
            padx=12,
            pady=6,
        )
        save_button.pack(side="right")

        self.window.protocol("WM_DELETE_WINDOW", lambda: self.save(close_after=True))
        self.update_style_buttons()
        self.textbox.focus_set()
        self.app.set_status("Type your note, then save and close the note window.")

    def configure_text_tags(self):
        if self.textbox is None:
            return
        self.textbox.tag_configure("bold", font=("Segoe UI", 10, "bold"))
        self.textbox.tag_configure("italic", font=("Segoe UI", 10, "italic"))
        self.textbox.tag_configure("bold_italic", font=("Segoe UI", 10, "bold italic"))

    def update_style_buttons(self):
        if self.window is None or not self.window.winfo_exists() or self.textbox is None:
            return
        has_selection = bool(self.textbox.tag_ranges("sel"))
        state = "normal" if has_selection else "disabled"
        self.bold_button.configure(bg="#e8edff", fg=TEXT_DARK, state=state)
        self.italic_button.configure(bg="#e8edff", fg=TEXT_DARK, state=state)

    def toggle_bold(self):
        self._toggle_selected_text_style("bold")

    def toggle_italic(self):
        self._toggle_selected_text_style("italic")

    def _toggle_selected_text_style(self, style_name):
        if self.textbox is None or not self.textbox.tag_ranges("sel"):
            self.app.set_status("Select text in the note box first.")
            self.update_style_buttons()
            return

        start = self.textbox.index("sel.first")
        end = self.textbox.index("sel.last")
        has_bold = self._style_range_has_tag("bold", start, end) or self._style_range_has_tag("bold_italic", start, end)
        has_italic = self._style_range_has_tag("italic", start, end) or self._style_range_has_tag("bold_italic", start, end)

        if style_name == "bold":
            wants_bold = not has_bold
            wants_italic = has_italic
        else:
            wants_bold = has_bold
            wants_italic = not has_italic

        for tag_name in ("bold", "italic", "bold_italic"):
            self.textbox.tag_remove(tag_name, start, end)

        if wants_bold and wants_italic:
            self.textbox.tag_add("bold_italic", start, end)
        elif wants_bold:
            self.textbox.tag_add("bold", start, end)
        elif wants_italic:
            self.textbox.tag_add("italic", start, end)

        self.update_style_buttons()
        self.textbox.focus_set()

    def _style_range_has_tag(self, tag_name, start, end):
        ranges = self.textbox.tag_ranges(tag_name)
        for i in range(0, len(ranges), 2):
            tag_start = ranges[i]
            tag_end = ranges[i + 1]
            if self.textbox.compare(tag_end, ">", start) and self.textbox.compare(tag_start, "<", end):
                return True
        return False

    def get_note_runs(self):
        if self.textbox is None or not self.window or not self.window.winfo_exists():
            return [(self.current_note_text, "normal")] if self.current_note_text else []

        end = self.textbox.index("end-1c")
        index = "1.0"
        runs = []
        current_style = None
        buffer = []

        while self.textbox.compare(index, "<", end):
            next_index = self.textbox.index(f"{index} +1c")
            char = self.textbox.get(index, next_index)
            tags = set(self.textbox.tag_names(index))

            if "bold_italic" in tags:
                style = "bold_italic"
            elif "bold" in tags:
                style = "bold"
            elif "italic" in tags:
                style = "italic"
            else:
                style = "normal"

            if current_style is None:
                current_style = style

            if style != current_style:
                runs.append(("".join(buffer), current_style))
                buffer = []
                current_style = style

            buffer.append(char)
            index = next_index

        if buffer:
            runs.append(("".join(buffer), current_style))

        return runs

    def append_pending_note_to_document(self):
        if not self.document_manager.has_document():
            return

        runs = self.get_note_runs()
        if not runs:
            return

        self.document_manager.add_note_runs(runs)
        self.document_manager.save()
        self.clear()


class PreviewManager:
    def __init__(self, app, document_manager):
        self.app = app
        self.document_manager = document_manager
        self.window = None
        self.canvas = None
        self.note_textbox = None
        self.note_save_after_id = None
        self.index_var = None
        self.buttons = {}
        self.paths = []
        self.current_index = 0
        self.current_image_path = None
        self.base_image = None
        self.annotated_image = None
        self.annotation_draw = None
        self.dirty = False
        self.zoom = 1.0
        self.pan_x = 0
        self.pan_y = 0
        self.tool_mode = "pan"
        self.brush_index = 1
        self.brush_size = PREVIEW_BRUSH_LEVELS[self.brush_index]
        self.display_size = (0, 0)
        self.preview_photo = None
        self.preview_drag_origin = None
        self.last_point = None
        self.loading_note = False

    def is_visible(self):
        return self.window is not None and self.window.winfo_exists() and self.window.state() != "withdrawn"

    def show(self):
        if self.window is not None and self.window.winfo_exists():
            self.window.deiconify()
            self.window.lift()
            self.window.focus_force()
            self.update_canvas()

    def hide(self):
        if self.window is not None and self.window.winfo_exists():
            self.window.withdraw()

    def refresh_paths(self):
        self.paths = list_saved_screenshots()
        if self.current_index >= len(self.paths):
            self.current_index = max(0, len(self.paths) - 1)
        return self.paths

    def _current_path(self):
        return self.paths[self.current_index] if self.paths else None

    def open_window(self):
        if self.window is not None and self.window.winfo_exists():
            self.show()
            return

        self.window = tk.Toplevel(self.app.root)
        self.window.title("Screenshot Preview")
        self.window.geometry("980x760+760+100")
        self.window.minsize(860, 680)
        self.window.configure(bg=CARD_BG)
        self.window.attributes("-topmost", True)

        top_bar = tk.Frame(self.window, bg=CARD_BG)
        top_bar.pack(fill="x", padx=14, pady=(14, 8))

        self.buttons["prev"] = tk.Button(
            top_bar,
            text="\u2190",
            command=lambda: self.show_at(self.current_index - 1),
            relief="flat",
            bg=CARD_BG,
            activebackground="#eef2ff",
            fg=TEXT_DARK,
            font=("Segoe UI Symbol", 13),
            bd=0,
            width=2,
            pady=1,
        )
        self.buttons["prev"].pack(side="left")

        self.index_var = tk.StringVar(value="0 / 0")
        tk.Label(
            top_bar,
            textvariable=self.index_var,
            bg=CARD_BG,
            fg=TEXT_DARK,
            font=("Segoe UI Semibold", 12),
        ).pack(side="left", padx=10)

        self.buttons["next"] = tk.Button(
            top_bar,
            text="\u2192",
            command=lambda: self.show_at(self.current_index + 1),
            relief="flat",
            bg=CARD_BG,
            activebackground="#eef2ff",
            fg=TEXT_DARK,
            font=("Segoe UI Symbol", 13),
            bd=0,
            width=2,
            pady=1,
        )
        self.buttons["next"].pack(side="left")

        self.buttons["zoom_out"] = tk.Button(
            top_bar,
            text="-",
            command=lambda: self.zoom_view(-0.25),
            relief="flat",
            bg=CARD_BG,
            activebackground="#eef2ff",
            fg=TEXT_DARK,
            font=("Segoe UI Semibold", 12),
            bd=0,
            width=2,
            pady=1,
        )
        self.buttons["zoom_out"].pack(side="left", padx=(16, 0))

        self.buttons["zoom_in"] = tk.Button(
            top_bar,
            text="+",
            command=lambda: self.zoom_view(0.25),
            relief="flat",
            bg=CARD_BG,
            activebackground="#eef2ff",
            fg=TEXT_DARK,
            font=("Segoe UI Semibold", 12),
            bd=0,
            width=2,
            pady=1,
        )
        self.buttons["zoom_in"].pack(side="left", padx=(8, 0))

        self.buttons["pan"] = tk.Button(
            top_bar,
            text="\u270b",
            command=lambda: self.set_tool_mode("pan"),
            relief="flat",
            bg=CARD_BG,
            activebackground="#eef2ff",
            fg=PRIMARY_BLUE,
            font=("Segoe UI Symbol", 12),
            bd=0,
            width=2,
            pady=1,
        )
        self.buttons["pan"].pack(side="left", padx=(16, 0))

        self.buttons["highlight"] = tk.Button(
            top_bar,
            text="\u270e",
            command=lambda: self.set_tool_mode("highlight"),
            relief="flat",
            bg=CARD_BG,
            activebackground="#fff6d8",
            fg=TEXT_DARK,
            font=("Segoe UI Symbol", 12),
            bd=0,
            width=2,
            pady=1,
        )
        self.buttons["highlight"].pack(side="left", padx=(8, 0))

        tk.Label(
            top_bar,
            bg=CARD_BG,
            fg="#6d7891",
            font=("Segoe UI Semibold", 9),
            text="Brush",
        ).pack(side="left", padx=(16, 4))

        for index, label in enumerate(("S", "M", "L", "XL")):
            self.buttons[f"brush_{label}"] = tk.Button(
                top_bar,
                text=label,
                command=lambda idx=index: self.set_brush_level(idx),
                relief="flat",
                bg=CARD_BG,
                activebackground="#fff6d8",
                fg=TEXT_DARK,
                font=("Segoe UI Semibold", 9),
                bd=0,
                width=3,
                pady=2,
            )
            self.buttons[f"brush_{label}"].pack(side="left", padx=(4 if index else 0, 0))

        self.buttons["delete"] = tk.Button(
            top_bar,
            text="\U0001F5D1",
            command=self.delete_current_image,
            relief="flat",
            bg=CARD_BG,
            activebackground="#ffecef",
            fg="#8b1e2d",
            font=("Segoe UI Symbol", 12),
            bd=0,
            width=2,
            pady=1,
        )
        self.buttons["delete"].pack(side="right")

        self.note_textbox = tk.Text(
            self.window,
            height=3,
            wrap="word",
            font=("Segoe UI", 10),
            bd=1,
            relief="solid",
            padx=10,
            pady=7,
            fg=TEXT_DARK,
            highlightthickness=1,
            highlightbackground="#d9deea",
        )
        self.note_textbox.pack(fill="x", padx=14, pady=(0, 10))
        self.note_textbox.bind("<KeyRelease>", self.schedule_note_save)
        self.note_textbox.bind("<FocusOut>", self.save_current_note)

        self.canvas = tk.Canvas(
            self.window,
            width=920,
            height=590,
            bg="#eff3ff",
            highlightthickness=1,
            highlightbackground="#d9deea",
            cursor="fleur",
        )
        self.canvas.pack(fill="both", expand=True, padx=14, pady=(0, 10))
        self.canvas.bind("<Configure>", lambda _event: self.update_canvas())
        self.canvas.bind("<ButtonPress-1>", self.start_interaction)
        self.canvas.bind("<B1-Motion>", self.drag_interaction)
        self.canvas.bind("<ButtonRelease-1>", self.end_interaction)

        action_bar = tk.Frame(self.window, bg=CARD_BG)
        action_bar.pack(fill="x", padx=14, pady=(0, 12))

        tk.Button(
            action_bar,
            text="Reset",
            command=self.reset_highlight,
            relief="flat",
            bg="#fff4c2",
            fg="#7a5b00",
            padx=10,
            pady=4,
        ).pack(side="left")

        self.window.protocol("WM_DELETE_WINDOW", self.close_window)
        self.set_tool_mode("pan")
        self._update_brush_buttons()
        self.show_at(len(self.refresh_paths()) - 1)

    def close_window(self):
        self.save_current_note()
        if self.window is not None and self.window.winfo_exists():
            self.window.destroy()
        self._reset_state()

    def _reset_state(self):
        if self.note_save_after_id is not None and self.window is not None and self.window.winfo_exists():
            self.window.after_cancel(self.note_save_after_id)
        self.window = None
        self.canvas = None
        self.note_textbox = None
        self.note_save_after_id = None
        self.index_var = None
        self.buttons.clear()
        self.paths = []
        self.current_index = 0
        self.current_image_path = None
        self.base_image = None
        self.annotated_image = None
        self.annotation_draw = None
        self.dirty = False
        self.zoom = 1.0
        self.pan_x = 0
        self.pan_y = 0
        self.tool_mode = "pan"
        self.brush_index = 1
        self.brush_size = PREVIEW_BRUSH_LEVELS[self.brush_index]
        self.display_size = (0, 0)
        self.preview_photo = None
        self.preview_drag_origin = None
        self.last_point = None
        self.loading_note = False

    def show_at(self, index):
        self.current_index = max(0, min(index, len(self.paths) - 1)) if self.paths else 0
        self.update_canvas()

    def zoom_view(self, delta):
        self.zoom = max(1.0, min(3.0, self.zoom + delta))
        if self.zoom == 1.0:
            self.pan_x = 0
            self.pan_y = 0
        self.update_canvas()

    def set_brush_level(self, index):
        self.brush_index = max(0, min(len(PREVIEW_BRUSH_LEVELS) - 1, index))
        self.brush_size = PREVIEW_BRUSH_LEVELS[self.brush_index]
        self._update_brush_buttons()

    def _update_brush_buttons(self):
        for index, label in enumerate(("S", "M", "L", "XL")):
            button = self.buttons.get(f"brush_{label}")
            if button is None or not button.winfo_exists():
                continue
            is_active = index == self.brush_index
            button.configure(
                bg="#ffe89a" if is_active else CARD_BG,
                fg="#7a5b00" if is_active else TEXT_DARK,
            )

    def set_tool_mode(self, mode):
        self.tool_mode = mode
        if self.canvas is not None and self.canvas.winfo_exists():
            self.canvas.configure(cursor="fleur" if mode == "pan" else "pencil")

        pan_button = self.buttons.get("pan")
        highlight_button = self.buttons.get("highlight")
        if pan_button is not None and pan_button.winfo_exists():
            pan_button.configure(fg=PRIMARY_BLUE if mode == "pan" else TEXT_DARK)
        if highlight_button is not None and highlight_button.winfo_exists():
            highlight_button.configure(fg="#b58900" if mode == "highlight" else TEXT_DARK)

    def start_interaction(self, event):
        self.preview_drag_origin = (event.x, event.y)
        if self.tool_mode == "highlight":
            self.last_point = self._canvas_to_image_coords(event)
        else:
            self.last_point = None

    def drag_interaction(self, event):
        if self.tool_mode == "highlight":
            if self.annotation_draw is None:
                return
            point = self._canvas_to_image_coords(event)
            if point is None:
                self.last_point = None
                return
            if self.last_point is None:
                self.last_point = point
                return
            draw = ImageDraw.Draw(self.annotation_draw)
            draw.line(
                (self.last_point[0], self.last_point[1], point[0], point[1]),
                fill=(255, 235, 59, 150),
                width=self.brush_size,
            )
            self.last_point = point
            self.dirty = True
            self.update_canvas()
            return

        if self.preview_drag_origin is None:
            self.preview_drag_origin = (event.x, event.y)
            return

        dx = event.x - self.preview_drag_origin[0]
        dy = event.y - self.preview_drag_origin[1]
        self.preview_drag_origin = (event.x, event.y)
        self.pan_x += dx
        self.pan_y += dy
        self.update_canvas()

    def end_interaction(self, _event):
        self.preview_drag_origin = None
        self.last_point = None
        if self.tool_mode == "highlight" and self.dirty:
            self.save_highlight()

    def update_canvas(self):
        if self.window is None or not self.window.winfo_exists() or self.canvas is None:
            return

        self.refresh_paths()
        if not self.paths:
            self.canvas.delete("all")
            self.canvas.create_text(
                190,
                150,
                text="No screenshots yet.\nTake one and it will appear here.",
                fill="#5d6880",
                font=("Segoe UI", 11),
                justify="center",
            )
            if self.index_var is not None:
                self.index_var.set("0 / 0")
            self.current_image_path = None
            self.base_image = None
            self.annotated_image = None
            self.annotation_draw = None
            self.dirty = False
            self.load_note_for_current_image()
            self._set_controls_enabled(False)
            return

        image_path = self._current_path()
        if image_path != self.current_image_path or self.base_image is None:
            self.save_current_note()
            base_image = Image.open(image_path).convert("RGBA")
            self.base_image = base_image
            self.annotated_image = base_image.copy()
            self.annotation_draw = Image.new("RGBA", base_image.size, (0, 0, 0, 0))
            self.current_image_path = image_path
            self.dirty = False
            self.pan_x = 0
            self.pan_y = 0
            self.load_note_for_current_image()

        canvas_width = max(240, self.canvas.winfo_width())
        canvas_height = max(180, self.canvas.winfo_height())
        display_image = Image.alpha_composite(self.annotated_image, self.annotation_draw)
        viewport_width = max(80, canvas_width - 16)
        viewport_height = max(80, canvas_height - 16)
        base_scale = min(viewport_width / display_image.width, viewport_height / display_image.height)
        render_scale = base_scale * self.zoom
        scaled_width = max(1, int(display_image.width * render_scale))
        scaled_height = max(1, int(display_image.height * render_scale))
        scaled_image = display_image.resize((scaled_width, scaled_height), Image.Resampling.LANCZOS)

        max_pan_x = max(0, (scaled_width - viewport_width) / 2)
        max_pan_y = max(0, (scaled_height - viewport_height) / 2)
        self.pan_x = max(-max_pan_x, min(max_pan_x, self.pan_x))
        self.pan_y = max(-max_pan_y, min(max_pan_y, self.pan_y))

        viewport = Image.new("RGBA", (viewport_width, viewport_height), (239, 243, 255, 255))
        paste_x = int((viewport_width - scaled_width) / 2 + self.pan_x)
        paste_y = int((viewport_height - scaled_height) / 2 + self.pan_y)
        viewport.paste(scaled_image, (paste_x, paste_y), scaled_image)

        self.display_size = (scaled_width, scaled_height)
        self.preview_photo = ImageTk.PhotoImage(viewport)

        self.canvas.delete("all")
        self.canvas.create_rectangle(0, 0, canvas_width, canvas_height, fill="#eff3ff", outline="")
        self.canvas.create_image(canvas_width / 2, canvas_height / 2, image=self.preview_photo)
        self.canvas.create_rectangle(8, 8, canvas_width - 8, canvas_height - 8, outline="#cfd7ef", width=1)

        if self.index_var is not None:
            self.index_var.set(f"{self.current_index + 1} / {len(self.paths)}")
        self._set_controls_enabled(True)

    def load_note_for_current_image(self):
        if self.note_textbox is None or not self.note_textbox.winfo_exists():
            return

        self.loading_note = True
        self.note_textbox.configure(state="normal")
        self.note_textbox.delete("1.0", tk.END)
        if self.current_image_path is not None:
            self.note_textbox.insert("1.0", get_preview_note(self.current_image_path))
        self.loading_note = False

    def schedule_note_save(self, _event=None):
        if self.loading_note or self.window is None or not self.window.winfo_exists():
            return
        if self.note_save_after_id is not None:
            self.window.after_cancel(self.note_save_after_id)
        self.note_save_after_id = self.window.after(800, self.save_current_note)

    def save_current_note(self, _event=None):
        if self.note_save_after_id is not None and self.window is not None and self.window.winfo_exists():
            self.window.after_cancel(self.note_save_after_id)
            self.note_save_after_id = None
        if self.note_textbox is None or not self.note_textbox.winfo_exists() or self.current_image_path is None:
            return

        note_text = self.note_textbox.get("1.0", tk.END).strip()
        set_preview_note(self.current_image_path, note_text)
        self.document_manager.set_image_note(self.current_image_path, note_text)

    def _set_controls_enabled(self, enabled):
        state = "normal" if enabled else "disabled"
        for key in (
            "prev",
            "next",
            "delete",
            "zoom_in",
            "zoom_out",
            "pan",
            "highlight",
            "brush_S",
            "brush_M",
            "brush_L",
            "brush_XL",
        ):
            button = self.buttons.get(key)
            if button is not None and button.winfo_exists():
                button.configure(state=state)
        if self.note_textbox is not None and self.note_textbox.winfo_exists():
            self.note_textbox.configure(state="normal" if enabled else "disabled")

    def _canvas_to_image_coords(self, event):
        if self.base_image is None or self.canvas is None or self.preview_photo is None:
            return None

        canvas_width = max(240, self.canvas.winfo_width())
        canvas_height = max(180, self.canvas.winfo_height())
        viewport_width = max(80, canvas_width - 16)
        viewport_height = max(80, canvas_height - 16)
        display_width, display_height = self.display_size
        origin_x = 8 + (viewport_width - display_width) / 2 + self.pan_x
        origin_y = 8 + (viewport_height - display_height) / 2 + self.pan_y

        if not (origin_x <= event.x <= origin_x + display_width and origin_y <= event.y <= origin_y + display_height):
            return None

        scale_x = self.base_image.width / display_width
        scale_y = self.base_image.height / display_height
        image_x = int((event.x - origin_x) * scale_x)
        image_y = int((event.y - origin_y) * scale_y)
        return image_x, image_y

    def reset_highlight(self):
        if self.current_image_path is None:
            self.app.set_status("Reset needs an original backup for this screenshot.")
            return

        backup_path = get_screenshot_backup_path(self.current_image_path)
        if os.path.exists(backup_path):
            restored_image = Image.open(backup_path).convert("RGB")
            restored_image.save(self.current_image_path)
            self.base_image = restored_image.convert("RGBA")
            self.annotated_image = self.base_image.copy()
            self.annotation_draw = Image.new("RGBA", self.base_image.size, (0, 0, 0, 0))
            self.dirty = False
            if self.document_manager.has_document():
                note_element = self.document_manager.tracked_notes.get(self.document_manager._normalize_path(self.current_image_path))
                self.document_manager.remove_tracked_image(self.current_image_path, remove_note=False)
                self.document_manager.append_image(self.current_image_path, width=Inches(5), note_element=note_element)
                self.document_manager.save()
            self.update_canvas()
            self.app.set_status(
                f"Reset screenshot to original in {os.path.basename(self.document_manager.path)}"
                if self.document_manager.has_document()
                else "Reset screenshot to original."
            )
            return

        if self.dirty:
            self.annotation_draw = Image.new("RGBA", self.base_image.size, (0, 0, 0, 0))
            self.dirty = False
            self.update_canvas()
            self.app.set_status("Unsaved highlight cleared.")
            return

        self.app.set_status("Reset needs an original backup for this screenshot.")

    def save_highlight(self):
        if not self.document_manager.has_document():
            messagebox.showwarning("No document", "Create or select a Word document first.")
            return

        if self.base_image is None or self.current_image_path is None:
            self.app.set_status("No screenshot selected in preview.")
            return

        if not self.dirty:
            self.app.set_status("Add a yellow highlight first, then save it.")
            return

        ensure_screenshot_backup(self.current_image_path)
        highlighted_image = Image.alpha_composite(self.annotated_image, self.annotation_draw).convert("RGB")
        highlighted_image.save(self.current_image_path)
        self.base_image = highlighted_image.convert("RGBA")
        self.annotated_image = self.base_image.copy()
        self.annotation_draw = Image.new("RGBA", self.base_image.size, (0, 0, 0, 0))
        note_element = self.document_manager.tracked_notes.get(self.document_manager._normalize_path(self.current_image_path))
        self.document_manager.remove_tracked_image(self.current_image_path, remove_note=False)
        self.document_manager.append_image(self.current_image_path, width=Inches(5), note_element=note_element)
        self.document_manager.save()
        self.dirty = False
        self.refresh_paths()
        self.update_canvas()
        self.app.set_status(f"Saved highlighted screenshot to {os.path.basename(self.document_manager.path)}")

    def delete_current_image(self):
        if self.current_image_path is None:
            self.app.set_status("No screenshot selected to delete.")
            return

        image_name = os.path.basename(self.current_image_path)
        backup_path = get_screenshot_backup_path(self.current_image_path)
        try:
            if self.document_manager.has_document():
                self.document_manager.remove_tracked_image(self.current_image_path)
                self.document_manager.save()
            os.remove(self.current_image_path)
            if os.path.exists(backup_path):
                os.remove(backup_path)
            remove_preview_note(self.current_image_path)
        except Exception as exc:
            messagebox.showerror("Delete failed", f"Could not delete screenshot.\n\n{exc}")
            return

        self.refresh_paths()
        if self.paths:
            self.show_at(min(self.current_index, len(self.paths) - 1))
        else:
            self.current_index = 0
            self.update_canvas()

        self.app.show_toast("Screenshot deleted")
        self.app.set_status(f"Deleted screenshot: {image_name}")


class DrawingManager:
    def __init__(self, app):
        self.app = app
        self.overlay = None
        self.toolbar = None
        self.canvas = None
        self.tool = "pencil"
        self.color = "#f61d2a"
        self.start_point = None
        self.last_point = None
        self.preview_item = None
        self.background_photo = None
        self.annotation_image = None
        self.annotation_draw = None
        self.tool_buttons = {}
        self.color_buttons = {}
        self.was_visible_for_capture = False

    def is_visible(self):
        return self.overlay is not None and self.overlay.winfo_exists() and self.overlay.state() != "withdrawn"

    def toggle(self):
        if self.is_visible():
            self.hide(clear=False)
            self.app.set_status("Draw hidden.")
        else:
            self.show()

    def show(self):
        self.app.root.withdraw()
        self.app.root.update()
        background = capture_image()

        if self.overlay is None or not self.overlay.winfo_exists():
            self._build_overlay(background)
        else:
            self._load_background(background)
        if self.toolbar is None or not self.toolbar.winfo_exists():
            self._build_toolbar()

        self.overlay.deiconify()
        self.overlay.lift()
        self.overlay.attributes("-topmost", True)
        self.overlay.focus_force()
        self.app.root.deiconify()
        self.app.root.lift()
        self.app.root.attributes("-topmost", True)
        self.toolbar.deiconify()
        self.toolbar.lift()
        self.toolbar.attributes("-topmost", True)
        self._bring_controls_to_front()
        self._apply_cursor()
        self._update_button_states()
        self.app.set_status("Draw mode on. Draw on the screen, then press Shot.")

    def hide(self, clear=False):
        if self.toolbar is not None and self.toolbar.winfo_exists():
            self.toolbar.withdraw()
        if self.overlay is not None and self.overlay.winfo_exists():
            self.overlay.withdraw()
        if clear and self.canvas is not None and self.canvas.winfo_exists():
            self.canvas.delete("all")
            self.annotation_image = None
            self.annotation_draw = None
            self.background_photo = None
        self.start_point = None
        self.last_point = None
        self.preview_item = None

    def _build_overlay(self, background):
        screen_width, screen_height = background.size
        self.overlay = tk.Toplevel(self.app.root)
        self.overlay.overrideredirect(True)
        self.overlay.attributes("-topmost", True)
        self.overlay.geometry(f"{screen_width}x{screen_height}+0+0")
        self.overlay.configure(bg="black")

        self.canvas = tk.Canvas(
            self.overlay,
            width=screen_width,
            height=screen_height,
            bg="black",
            highlightthickness=0,
            cursor="crosshair",
        )
        self.canvas.pack(fill="both", expand=True)
        self.canvas.bind("<ButtonPress-1>", self.start_draw)
        self.canvas.bind("<B1-Motion>", self.drag_draw)
        self.canvas.bind("<ButtonRelease-1>", self.end_draw)
        self.overlay.bind("<Escape>", lambda _event: self.hide(clear=False))
        self._load_background(background)

    def _load_background(self, background):
        if self.canvas is None or not self.canvas.winfo_exists():
            return
        self.annotation_image = Image.new("RGBA", background.size, (0, 0, 0, 0))
        self.annotation_draw = ImageDraw.Draw(self.annotation_image)
        self.background_photo = ImageTk.PhotoImage(background)
        self.canvas.delete("all")
        self.canvas.config(width=background.width, height=background.height)
        self.canvas.create_image(0, 0, image=self.background_photo, anchor="nw", tags=("background",))
        self.canvas.tag_lower("background")

    def _build_toolbar(self):
        self.toolbar = tk.Toplevel(self.app.root)
        self.toolbar.overrideredirect(True)
        self.toolbar.attributes("-topmost", True)
        self.toolbar.configure(bg=CARD_BG)
        self.toolbar.geometry("+80+80")

        frame = tk.Frame(self.toolbar, bg=CARD_BG, bd=1, relief="solid")
        frame.pack()

        tool_specs = (
            ("pencil", "\u270e"),
            ("line", "/"),
            ("rect", "\u25a1"),
            ("circle", "\u25cb"),
        )
        for tool_name, icon in tool_specs:
            button = tk.Button(
                frame,
                text=icon,
                width=2,
                relief="flat",
                bd=0,
                bg=CARD_BG,
                fg=TEXT_DARK,
                activebackground="#eef2ff",
                command=lambda name=tool_name: self.set_tool(name),
            )
            button.pack(side="left", padx=2, pady=3)
            self.tool_buttons[tool_name] = button

        for color in ("#f61d2a", "#f7d038", "#188038", "#1a73e8", "#111827"):
            button = tk.Button(
                frame,
                width=2,
                relief="flat",
                bd=0,
                bg=color,
                activebackground=color,
                command=lambda value=color: self.set_color(value),
            )
            button.pack(side="left", padx=2, pady=3)
            self.color_buttons[color] = button

    def set_tool(self, tool):
        self.tool = tool
        self._apply_cursor()
        self._update_button_states()
        self._bring_controls_to_front()

    def set_color(self, color):
        self.color = color
        self._update_button_states()
        self._bring_controls_to_front()

    def _update_button_states(self):
        for tool, button in self.tool_buttons.items():
            if button.winfo_exists():
                button.configure(bg="#e8edff" if tool == self.tool else CARD_BG)
        for color, button in self.color_buttons.items():
            if button.winfo_exists():
                button.configure(relief="solid" if color == self.color else "flat", bd=2 if color == self.color else 0)

    def _apply_cursor(self):
        if self.canvas is None or not self.canvas.winfo_exists():
            return
        cursor_by_tool = {
            "pencil": "pencil",
            "line": "crosshair",
            "rect": "crosshair",
            "circle": "crosshair",
        }
        self.canvas.configure(cursor=cursor_by_tool.get(self.tool, "crosshair"))

    def _bring_controls_to_front(self):
        if self.overlay is not None and self.overlay.winfo_exists():
            self.overlay.attributes("-topmost", True)
            try:
                self.overlay.lower(self.app.root)
            except tk.TclError:
                pass
        if self.app.root is not None and self.app.root.winfo_exists():
            self.app.root.deiconify()
            self.app.root.attributes("-topmost", True)
            self.app.root.lift()
        if self.toolbar is not None and self.toolbar.winfo_exists():
            self.toolbar.deiconify()
            self.toolbar.attributes("-topmost", True)
            self.toolbar.lift()

    def start_draw(self, event):
        self.start_point = (event.x, event.y)
        self.last_point = (event.x, event.y)
        self.preview_item = None

    def drag_draw(self, event):
        if self.canvas is None or self.start_point is None:
            return

        if self.tool == "pencil":
            self.canvas.create_line(
                self.last_point[0],
                self.last_point[1],
                event.x,
                event.y,
                fill=self.color,
                width=3,
                capstyle=tk.ROUND,
                smooth=True,
            )
            if self.annotation_draw is not None:
                self.annotation_draw.line(
                    (self.last_point[0], self.last_point[1], event.x, event.y),
                    fill=self.color,
                    width=3,
                )
            self.last_point = (event.x, event.y)
            return

        if self.preview_item is not None:
            self.canvas.delete(self.preview_item)

        x1, y1 = self.start_point
        x2, y2 = event.x, event.y
        if self.tool == "line":
            self.preview_item = self.canvas.create_line(x1, y1, x2, y2, fill=self.color, width=3)
        elif self.tool == "rect":
            if event.state & 0x0001:
                x2, y2 = self._square_endpoint(x1, y1, x2, y2)
            self.preview_item = self.canvas.create_rectangle(x1, y1, x2, y2, outline=self.color, width=3)
        elif self.tool == "circle":
            x2, y2 = self._square_endpoint(x1, y1, x2, y2)
            self.preview_item = self.canvas.create_oval(x1, y1, x2, y2, outline=self.color, width=3)

    def end_draw(self, event):
        if self.tool != "pencil":
            self.drag_draw(event)
            if self.annotation_draw is not None and self.start_point is not None:
                x1, y1 = self.start_point
                x2, y2 = event.x, event.y
                if self.tool == "line":
                    self.annotation_draw.line((x1, y1, x2, y2), fill=self.color, width=3)
                elif self.tool == "rect":
                    if event.state & 0x0001:
                        x2, y2 = self._square_endpoint(x1, y1, x2, y2)
                    self.annotation_draw.rectangle(self._ordered_bbox(x1, y1, x2, y2), outline=self.color, width=3)
                elif self.tool == "circle":
                    x2, y2 = self._square_endpoint(x1, y1, x2, y2)
                    self.annotation_draw.ellipse(self._ordered_bbox(x1, y1, x2, y2), outline=self.color, width=3)
        self.start_point = None
        self.last_point = None
        self.preview_item = None
        self._bring_controls_to_front()
        if self.overlay is not None and self.overlay.winfo_exists():
            self.overlay.after(80, self._bring_controls_to_front)

    def _square_endpoint(self, x1, y1, x2, y2):
        side = max(abs(x2 - x1), abs(y2 - y1))
        x2 = x1 + side if x2 >= x1 else x1 - side
        y2 = y1 + side if y2 >= y1 else y1 - side
        return x2, y2

    def _ordered_bbox(self, x1, y1, x2, y2):
        left, right = sorted((x1, x2))
        top, bottom = sorted((y1, y2))
        return left, top, right, bottom

    def prepare_for_capture(self):
        self.was_visible_for_capture = self.is_visible()
        if self.toolbar is not None and self.toolbar.winfo_exists():
            self.toolbar.withdraw()
        if self.overlay is not None and self.overlay.winfo_exists():
            self.overlay.withdraw()
        self.app.root.update()

    def apply_to_screenshot(self, screenshot, exclude_taskbar=False):
        if self.annotation_image is None:
            return screenshot

        annotation = self.annotation_image
        if exclude_taskbar:
            left, top, right, bottom = get_workarea_bbox()
            annotation = annotation.crop((left, top, right, bottom))

        base = screenshot.convert("RGBA")
        if annotation.size != base.size:
            annotation = annotation.resize(base.size, Image.Resampling.LANCZOS)
        return Image.alpha_composite(base, annotation).convert("RGB")

    def finish_capture(self):
        if self.was_visible_for_capture:
            self.hide(clear=True)
            self.was_visible_for_capture = False


class HotkeyManager:
    def __init__(self, app):
        self.app = app
        self.thread = None
        self.thread_id = None
        self.hotkey_registered = False
        self.workarea_hotkey_registered = False

    def start(self):
        if os.name != "nt":
            return
        self.thread = threading.Thread(target=self._hotkey_listener_loop, daemon=True)
        self.thread.start()

    def _hotkey_listener_loop(self):
        user32 = ctypes.windll.user32
        kernel32 = ctypes.windll.kernel32
        self.thread_id = kernel32.GetCurrentThreadId()

        if user32.RegisterHotKey(None, HOTKEY_ID, MOD_CONTROL | MOD_ALT, VK_Q):
            self.hotkey_registered = True
        else:
            self.app.schedule(self.app.set_status, f"Global shortcut unavailable: {SHOT_SHORTCUT_LABEL}")

        if user32.RegisterHotKey(None, WORKAREA_HOTKEY_ID, MOD_CONTROL | MOD_ALT, VK_W):
            self.workarea_hotkey_registered = True
        else:
            self.app.schedule(self.app.set_status, f"Taskbar-free shortcut unavailable: {WORKAREA_SHORTCUT_LABEL}")

        if not self.hotkey_registered and not self.workarea_hotkey_registered:
            return

        msg = wintypes.MSG()
        while user32.GetMessageW(ctypes.byref(msg), None, 0, 0) != 0:
            if msg.message == WM_HOTKEY:
                if msg.wParam == HOTKEY_ID:
                    self.app.schedule(self.app.take_screenshot)
                elif msg.wParam == WORKAREA_HOTKEY_ID:
                    self.app.schedule(self.app.take_screenshot, exclude_taskbar=True)
            user32.TranslateMessage(ctypes.byref(msg))
            user32.DispatchMessageW(ctypes.byref(msg))

        self._unregister_hotkeys(user32)

    def _unregister_hotkeys(self, user32):
        if self.hotkey_registered:
            user32.UnregisterHotKey(None, HOTKEY_ID)
            self.hotkey_registered = False
        if self.workarea_hotkey_registered:
            user32.UnregisterHotKey(None, WORKAREA_HOTKEY_ID)
            self.workarea_hotkey_registered = False

    def stop(self):
        if os.name != "nt":
            return

        user32 = ctypes.windll.user32
        self._unregister_hotkeys(user32)

        if self.thread_id:
            user32.PostThreadMessageW(self.thread_id, WM_QUIT, 0, 0)
            self.thread_id = None


class ButtonClickTracker:
    def __init__(self, root, single_command, double_command=None):
        self.root = root
        self.single_command = single_command
        self.double_command = double_command
        self.pending_id = None

    def bind(self, canvas, item):
        if self.double_command is None:
            canvas.tag_bind(item, "<Button-1>", self._single_click)
        else:
            canvas.tag_bind(item, "<Button-1>", self._schedule_single_click)
            canvas.tag_bind(item, "<Double-Button-1>", self._double_click)

    def _single_click(self, event):
        self.single_command()
        return "break"

    def _schedule_single_click(self, event):
        if self.pending_id is not None:
            self.root.after_cancel(self.pending_id)
        self.pending_id = self.root.after(SHOT_SINGLE_CLICK_DELAY_MS, self._execute_single)
        return "break"

    def _execute_single(self):
        self.pending_id = None
        self.single_command()

    def _double_click(self, event):
        if self.pending_id is not None:
            self.root.after_cancel(self.pending_id)
            self.pending_id = None
        self.double_command()
        return "break"


class App:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("Tester Tool")
        self.ui_scale = 0.80
        self.status_var = tk.StringVar(value="")
        self.canvas = None
        self.status_card = None
        self.status_label = None
        self.size_window = None
        self.size_slider = None
        self.tooltip_window = None
        self.toast_window = None
        self.toast_after_id = None

        self.document_manager = DocumentManager(self)
        self.note_manager = NoteManager(self, self.document_manager)
        self.preview_manager = PreviewManager(self, self.document_manager)
        self.drawing_manager = DrawingManager(self)
        self.hotkey_manager = HotkeyManager(self)
        self._main_thread_actions = queue.Queue()

        self._build_root_window()
        self.render_main_ui()
        self._process_main_thread_actions()
        self.hotkey_manager.start()

    def _build_root_window(self):
        self.root.geometry(f"{self.scaled(BASE_WIDTH)}x{self.scaled(BASE_HEIGHT)}+500+120")
        self.root.configure(bg=TRANSPARENT_KEY)
        self.root.attributes("-topmost", True)
        self.root.overrideredirect(True)
        try:
            self.root.wm_attributes("-transparentcolor", TRANSPARENT_KEY)
        except tk.TclError:
            pass

        self.canvas = tk.Canvas(self.root, width=420, height=680, bg=TRANSPARENT_KEY, highlightthickness=0)
        self.canvas.pack(fill="both", expand=True)

    def scaled(self, value):
        return max(1, int(value * self.ui_scale))

    def take_region_screenshot(self):
        self.take_screenshot(use_region_selector=True)

    def schedule(self, callback, *args, **kwargs):
        self._main_thread_actions.put((callback, args, kwargs))

    def _process_main_thread_actions(self):
        while True:
            try:
                callback, args, kwargs = self._main_thread_actions.get_nowait()
            except queue.Empty:
                break
            try:
                callback(*args, **kwargs)
            except Exception:
                pass
        self.root.after(50, self._process_main_thread_actions)

    def render_main_ui(self):
        current_x = self.root.winfo_x() if self.root.winfo_ismapped() else 500
        current_y = self.root.winfo_y() if self.root.winfo_ismapped() else 120
        self.root.geometry(f"{self.scaled(BASE_WIDTH)}x{self.scaled(BASE_HEIGHT)}+{current_x}+{current_y}")

        self.canvas.config(width=self.scaled(BASE_WIDTH), height=self.scaled(BASE_HEIGHT))
        self.canvas.delete("all")
        self._draw_background()
        if self.status_card is not None and self.status_card.winfo_exists():
            self.status_card.destroy()

        center_x = 182
        center_y = 318
        shot_size = 132
        satellite_size = 72
        shot_left = center_x - (shot_size / 2)
        shot_top = center_y - (shot_size / 2)
        self._draw_circle_button(
            shot_left,
            shot_top,
            shot_size,
            PRIMARY_RED,
            "Shot",
            self.take_screenshot,
            "",
            label_size=17,
            double_command=self.take_region_screenshot,
        )

        radius = 144
        button_specs = [
            ("Note", self.note_manager.open_window, "Open the note window.", 0, 10, False),
            ("Preview", self.preview_manager.open_window, "Browse earlier screenshots and add yellow highlights.", 45, 8, False),
            ("Draw", self.drawing_manager.toggle, "Show drawing tools.", 90, 9, False),
            ("i", self.open_info_window, "Show instructions.", 135, 18, False),
            ("Close", self.close_app, "Close the floating tool.", 180, 9, False),
            ("Move", None, "Drag this button to move the floating tool.", 225, 8, True),
            ("New", self.create_new_document, "Create a new Word document.", 270, 10, False),
            ("File", self.document_manager.select_document, "Open an existing Word document.", 315, 10, False),
        ]

        for label, command, tooltip_text, angle_deg, label_size, is_move in button_specs:
            radians = math.radians(angle_deg)
            button_center_x = center_x + radius * math.cos(radians)
            button_center_y = center_y + radius * math.sin(radians)
            button_left = button_center_x - (satellite_size / 2)
            button_top = button_center_y - (satellite_size / 2)

            if is_move:
                self._draw_move_button(button_left, button_top, satellite_size, PRIMARY_BLUE, label, tooltip_text, label_size=label_size)
            else:
                self._draw_circle_button(button_left, button_top, satellite_size, PRIMARY_BLUE, label, command, tooltip_text, label_size=label_size)

    def _draw_background(self):
        self.canvas.create_rectangle(
            0,
            0,
            self.scaled(BASE_WIDTH),
            self.scaled(BASE_HEIGHT),
            fill=TRANSPARENT_KEY,
            outline=TRANSPARENT_KEY,
        )

    def _draw_circle_button(self, x, y, size, fill, label, command, tooltip_text, label_size=11, double_command=None):
        x = self.scaled(x)
        y = self.scaled(y)
        size = self.scaled(size)
        label_size = self.scaled(label_size)
        shadow_offset = 6
        shadow = self.canvas.create_oval(
            x + self.scaled(shadow_offset),
            y + self.scaled(shadow_offset),
            x + size + self.scaled(shadow_offset),
            y + size + self.scaled(shadow_offset),
            fill=BUTTON_SHADOW,
            outline="",
        )
        oval = self.canvas.create_oval(x, y, x + size, y + size, fill=fill, outline="")
        text = self.canvas.create_text(
            x + size / 2,
            y + size / 2,
            text=label,
            fill="white",
            font=("Segoe UI Semibold", label_size),
            justify="center",
        )

        tracker = ButtonClickTracker(self.root, command, double_command=double_command)
        tracker.bind(self.canvas, shadow)
        tracker.bind(self.canvas, oval)
        tracker.bind(self.canvas, text)
        self._bind_tooltip((shadow, oval, text), tooltip_text)

    def _draw_move_button(self, x, y, size, fill, label, tooltip_text, label_size=11):
        x = self.scaled(x)
        y = self.scaled(y)
        size = self.scaled(size)
        label_size = self.scaled(label_size)
        shadow_offset = 6
        shadow = self.canvas.create_oval(
            x + self.scaled(shadow_offset),
            y + self.scaled(shadow_offset),
            x + size + self.scaled(shadow_offset),
            y + size + self.scaled(shadow_offset),
            fill=BUTTON_SHADOW,
            outline="",
        )
        oval = self.canvas.create_oval(x, y, x + size, y + size, fill=fill, outline="")
        text = self.canvas.create_text(
            x + size / 2,
            y + size / 2,
            text=label,
            fill="white",
            font=("Segoe UI Semibold", label_size),
            justify="center",
        )

        for item in (shadow, oval, text):
            self.canvas.tag_bind(item, "<Button-1>", self.start_move)
            self.canvas.tag_bind(item, "<B1-Motion>", self.on_motion)
            self.canvas.tag_bind(item, "<Button-3>", lambda event: self.open_size_window(event))
        self._bind_tooltip((shadow, oval, text), tooltip_text)

    def _bind_tooltip(self, items, tooltip_text):
        if not tooltip_text:
            return
        for item in items:
            self.canvas.tag_bind(
                item,
                "<Enter>",
                lambda event, text=tooltip_text: self.show_tooltip(text, event.x_root, event.y_root),
            )
            self.canvas.tag_bind(item, "<Leave>", lambda _event: self.hide_tooltip())

    def start_move(self, event):
        self.root.x = event.x_root - self.root.winfo_x()
        self.root.y = event.y_root - self.root.winfo_y()
        return "break"

    def on_motion(self, event):
        x = event.x_root - self.root.x
        y = event.y_root - self.root.y
        self.root.geometry(f"+{x}+{y}")
        return "break"

    def open_size_window(self, _event=None):
        if self.size_window is not None and self.size_window.winfo_exists():
            self.size_window.deiconify()
            self.size_window.lift()
            self.size_window.focus_force()
            return

        self.size_window = tk.Toplevel(self.root)
        self.size_window.title("Resize Tool")
        self.size_window.geometry(f"320x140+{self.root.winfo_x() + self.scaled(460)}+{self.root.winfo_y() + 40}")
        self.size_window.configure(bg=CARD_BG)
        self.size_window.attributes("-topmost", True)

        tk.Label(
            self.size_window,
            text="Adjust floating tool size",
            bg=CARD_BG,
            fg=TEXT_DARK,
            font=("Segoe UI Semibold", 11),
        ).pack(anchor="w", padx=16, pady=(14, 8))

        self.size_slider = tk.Scale(
            self.size_window,
            from_=int(MIN_SCALE * 100),
            to=int(MAX_SCALE * 100),
            orient="horizontal",
            command=self.apply_scale_from_slider,
            bg=CARD_BG,
            fg=TEXT_DARK,
            highlightthickness=0,
            length=260,
        )
        self.size_slider.pack(padx=16, fill="x")
        self.size_slider.set(int(self.ui_scale * 100))

        tk.Label(
            self.size_window,
            text="Tip: right-click the floating tool anytime to resize it.",
            bg=CARD_BG,
            fg="#5d6880",
            font=("Segoe UI", 9),
        ).pack(anchor="w", padx=16, pady=(6, 10))

        self.size_window.protocol("WM_DELETE_WINDOW", self.close_size_window)

    def close_size_window(self):
        if self.size_window is not None and self.size_window.winfo_exists():
            self.size_window.destroy()
        self.size_window = None
        self.size_slider = None

    def apply_scale_from_slider(self, value):
        self.apply_ui_scale(float(value))

    def apply_ui_scale(self, new_scale):
        self.ui_scale = max(MIN_SCALE, min(MAX_SCALE, new_scale))
        if self.size_slider is not None and self.size_slider.winfo_exists():
            current = int(float(self.size_slider.get()))
            target = int(self.ui_scale * 100)
            if current != target:
                self.size_slider.set(target)
        self.render_main_ui()

    def show_tooltip(self, text, x, y):
        self.hide_tooltip()
        self.tooltip_window = tk.Toplevel(self.root)
        self.tooltip_window.overrideredirect(True)
        self.tooltip_window.attributes("-topmost", True)
        self.tooltip_window.configure(bg="#182235")
        self.tooltip_window.geometry(f"+{x + 14}+{y + 10}")

        label = tk.Label(
            self.tooltip_window,
            text=text,
            bg="#182235",
            fg="white",
            font=("Segoe UI", 9),
            padx=10,
            pady=6,
            justify="left",
        )
        label.pack()

    def hide_tooltip(self):
        if self.tooltip_window is not None and self.tooltip_window.winfo_exists():
            self.tooltip_window.destroy()
        self.tooltip_window = None

    def show_toast(self, message, duration_ms=3000):
        self.close_toast()
        self.toast_window = tk.Toplevel(self.root)
        self.toast_window.overrideredirect(True)
        self.toast_window.attributes("-topmost", True)
        self.toast_window.configure(bg="#188038")

        label = tk.Label(
            self.toast_window,
            text=message,
            bg="#188038",
            fg="white",
            font=("Segoe UI Semibold", 12),
            padx=22,
            pady=14,
        )
        label.pack()

        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()
        self.toast_window.update_idletasks()
        width = self.toast_window.winfo_width()
        height = self.toast_window.winfo_height()
        x = int((screen_width - width) / 2)
        y = max(24, int(screen_height * 0.08))
        self.toast_window.geometry(f"+{x}+{y}")

        self.toast_after_id = self.root.after(duration_ms, self.close_toast)

    def close_toast(self):
        if self.toast_after_id is not None:
            self.root.after_cancel(self.toast_after_id)
            self.toast_after_id = None
        if self.toast_window is not None and self.toast_window.winfo_exists():
            self.toast_window.destroy()
        self.toast_window = None

    def create_new_document(self):
        folder = filedialog.askdirectory(title="Select folder to save the new document")
        if not folder:
            self.set_status("New document cancelled.")
            return

        title_window = tk.Toplevel(self.root)
        title_window.title("New Document Title")
        title_window.geometry("480x240+1020+160")
        title_window.configure(bg=CARD_BG)
        title_window.attributes("-topmost", True)

        tk.Label(
            title_window,
            text="Enter document title",
            bg=CARD_BG,
            fg=TEXT_DARK,
            font=("Segoe UI Semibold", 12),
        ).pack(anchor="w", padx=18, pady=(16, 8))

        tk.Label(
            title_window,
            text=f"Folder: {folder}",
            bg=CARD_BG,
            fg="#5d6880",
            font=("Segoe UI", 9),
            wraplength=430,
            justify="left",
        ).pack(anchor="w", padx=18, pady=(0, 8))

        title_textbox = tk.Text(
            title_window,
            height=4,
            width=40,
            wrap="word",
            font=("Segoe UI", 12),
            bd=1,
            relief="solid",
            padx=12,
            pady=12,
            fg=TEXT_DARK,
            highlightthickness=1,
            highlightbackground="#d9deea",
        )
        title_textbox.pack(fill="x", padx=18, pady=(0, 14))
        title_textbox.insert("1.0", datetime.now().strftime("TestDoc_%Y%m%d_%H%M%S"))
        title_textbox.focus_set()

        def save_new_document(_event=None):
            title = title_textbox.get("1.0", tk.END).strip()
            if not title:
                self.set_status("Enter a title for the new document.")
                return "break"

            if self.document_manager.create_new_document_file(folder, title):
                title_window.destroy()
            return "break"

        title_textbox.bind("<Return>", save_new_document)

        action_bar = tk.Frame(title_window, bg=CARD_BG)
        action_bar.pack(fill="x", padx=18, pady=(0, 16))

        tk.Button(
            action_bar,
            text="Create",
            command=save_new_document,
            bg=PRIMARY_BLUE,
            fg="white",
            relief="flat",
            padx=14,
            pady=6,
        ).pack(side="right")

        title_window.protocol("WM_DELETE_WINDOW", title_window.destroy)

    def open_info_window(self):
        info_window = tk.Toplevel(self.root)
        info_window.title("How To Use")
        info_window.geometry("500x440+1040+140")
        info_window.configure(bg=CARD_BG)
        info_window.attributes("-topmost", True)

        tk.Label(
            info_window,
            text="How to use",
            bg=CARD_BG,
            fg=TEXT_DARK,
            font=("Segoe UI Semibold", 12),
        ).pack(anchor="w", padx=18, pady=(16, 8))

        content_frame = tk.Frame(info_window, bg=CARD_BG)
        content_frame.pack(fill="both", expand=True, padx=18, pady=(0, 12))

        scrollbar = tk.Scrollbar(content_frame)
        scrollbar.pack(side="right", fill="y")

        info_text = tk.Text(
            content_frame,
            bg=CARD_BG,
            fg=TEXT_DARK,
            font=("Segoe UI", 10),
            wrap="word",
            relief="flat",
            bd=0,
            padx=4,
            pady=4,
            yscrollcommand=scrollbar.set,
        )
        info_text.pack(side="left", fill="both", expand=True)
        scrollbar.config(command=info_text.yview)

        instructions = (
            "How this tool works\n\n"
            "New button\n"
            "- Opens folder selection first.\n"
            "- Then shows a larger title box where you can enter the document title.\n"
            "- After clicking Create, a new Word document is created in that folder.\n\n"
            "File button\n"
            "- Opens an existing Word document.\n"
            "- The selected file becomes the active document for future screenshots.\n\n"
            "Note button\n"
            "- Opens the note window.\n"
            "- You can type text, use Bold or Italic, and then click Save & Close.\n"
            "- The saved note is attached to the next screenshot entry in the document.\n\n"
            "Shot button\n"
            "- Hides the floating tool, preview window, and note window before capture.\n"
            "- Takes the screenshot with the taskbar visible.\n"
            "- Double-click Shot to drag and capture only the required screen region.\n"
            "- Adds the note text first, then the screenshot image into the active document.\n"
            "- Saves the image into the screenshots folder also.\n"
            "- Clears the note after the capture is saved.\n"
            f"- Keyboard shortcut: {SHOT_SHORTCUT_LABEL}\n"
            f"- Taskbar-free shortcut: {WORKAREA_SHORTCUT_LABEL}\n\n"
            "Preview button\n"
            "- Opens a compact screenshot preview window.\n"
            "- Lets you move through saved screenshots one by one.\n"
            "- You can drag a yellow marker on top of the selected screenshot.\n"
            "- Save Highlight updates the selected screenshot instead of creating a duplicate copy.\n"
            "- Delete removes the selected screenshot from the preview list and screenshots folder.\n\n"
            "i button\n"
            "- Opens this help window.\n"
            "- You can scroll here to read all instructions.\n\n"
            "Close button\n"
            "- Before closing, if any note text is still pending, it is added at the end of the active document.\n"
            "- All temporary images in the screenshots folder are deleted after the save step completes.\n"
            "- Then the floating tool closes.\n\n"
            "General flow\n"
            "1. Click New or File.\n"
            "2. Add your note if needed.\n"
            "3. Click Shot.\n"
            "4. Repeat as many times as needed.\n\n"
            "By Tester/Developer\n"
            "Kalkeesh Jami\n"
            "email: kalkeesh.jami@accenture.com\n"
            "ph no: 7702726236\n"
            "lovely regards"
        )
        info_text.insert("1.0", instructions)
        info_text.configure(state="disabled")

        tk.Button(
            info_window,
            text="Close",
            command=info_window.destroy,
            bg=PRIMARY_BLUE,
            fg="white",
            relief="flat",
            padx=14,
            pady=6,
        ).pack(anchor="e", padx=18, pady=(0, 16))

    def take_screenshot(self, exclude_taskbar=False, use_region_selector=False):
        if not self.document_manager.has_document():
            messagebox.showwarning("No document", "Create or select a Word document first.")
            return

        hidden_windows = self._hide_autodoc_windows_for_capture()
        self.root.after(350, lambda: self._capture_after_hide(exclude_taskbar, use_region_selector, hidden_windows))

    def _hide_autodoc_windows_for_capture(self):
        if self.note_manager.is_visible():
            self.note_manager.save(close_after=False)
        if self.preview_manager.is_visible():
            self.preview_manager.save_current_note()
        self.drawing_manager.prepare_for_capture()

        hidden_windows = []
        for window in self.root.winfo_children():
            if not isinstance(window, tk.Toplevel):
                continue
            if window is self.drawing_manager.overlay:
                continue
            try:
                if window.winfo_exists() and window.state() != "withdrawn":
                    hidden_windows.append(window)
                    window.withdraw()
            except tk.TclError:
                continue

        self.hide_tooltip()
        self.close_toast()
        self.root.withdraw()
        self.root.update()
        return hidden_windows

    def _capture_after_hide(self, exclude_taskbar, use_region_selector, hidden_windows):
        note_runs = self.note_manager.get_note_runs()
        screenshot_dir = ensure_screenshot_dir()
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        suffix = "_notaskbar" if exclude_taskbar else ""
        img_path = os.path.join(screenshot_dir, f"screenshot_{timestamp}{suffix}.png")

        try:
            screenshot = capture_image(exclude_taskbar=exclude_taskbar)
            screenshot = self.drawing_manager.apply_to_screenshot(screenshot, exclude_taskbar=exclude_taskbar)
            if use_region_selector:
                region = self.select_capture_region(screenshot)
                if region is None:
                    self._restore_after_capture(hidden_windows)
                    self.set_status("Region capture cancelled.")
                    return
                screenshot = screenshot.crop(region)

            screenshot.save(img_path)
            ensure_screenshot_backup(img_path)

            note_element = None
            note_text = "".join(text for text, _style in note_runs).strip()
            if note_runs:
                note_element = self.document_manager.add_note_runs(note_runs)

            self.document_manager.append_image(img_path, width=Inches(5), note_element=note_element)
            set_preview_note(img_path, note_text)
            self.document_manager.save()
        except Exception as exc:
            messagebox.showerror("Capture failed", f"Could not save screenshot.\n\n{exc}")
            self.set_status("Screenshot failed.")
        else:
            self.note_manager.clear()
            self.preview_manager.refresh_paths()
            if self.preview_manager.is_visible():
                self.preview_manager.show_at(len(self.preview_manager.paths) - 1)
            self.show_toast("Screenshot added")
            capture_mode = "selected region" if use_region_selector else ("without taskbar" if exclude_taskbar else "with taskbar")
            self.set_status(f"Saved screenshot {capture_mode} to {img_path}")
        finally:
            self.drawing_manager.finish_capture()
            self._restore_after_capture(hidden_windows)

    def _restore_after_capture(self, hidden_windows):
        self.root.deiconify()
        for window in hidden_windows:
            try:
                if window.winfo_exists():
                    window.deiconify()
            except tk.TclError:
                continue
        if self.note_manager.is_visible():
            self.note_manager.show()
        if self.preview_manager.is_visible():
            self.preview_manager.show()

    def select_capture_region(self, base_image):
        selection = {"start": None, "end": None, "result": None}
        overlay = tk.Toplevel(self.root)
        overlay.overrideredirect(True)
        overlay.attributes("-topmost", True)
        overlay.geometry(f"{base_image.width}x{base_image.height}+0+0")
        overlay.configure(bg="black")

        overlay_photo = ImageTk.PhotoImage(base_image)
        canvas_overlay = tk.Canvas(overlay, width=base_image.width, height=base_image.height, highlightthickness=0, cursor="crosshair")
        canvas_overlay.pack(fill="both", expand=True)
        canvas_overlay.create_image(0, 0, image=overlay_photo, anchor="nw")
        shade = canvas_overlay.create_rectangle(0, 0, base_image.width, base_image.height, fill="black", stipple="gray50", outline="")
        rect = canvas_overlay.create_rectangle(0, 0, 0, 0, outline="#f7d038", width=2)
        hint = canvas_overlay.create_text(
            24,
            24,
            anchor="nw",
            text="Drag to capture a region. Press Esc to cancel.",
            fill="white",
            font=("Segoe UI Semibold", 12),
        )
        size_hint = canvas_overlay.create_text(
            24,
            52,
            anchor="nw",
            text="",
            fill="#ffe89a",
            font=("Segoe UI Semibold", 10),
        )
        canvas_overlay.tag_raise(rect)
        canvas_overlay.tag_raise(hint)
        canvas_overlay.tag_raise(size_hint)

        def on_press(event):
            selection["start"] = (event.x, event.y)
            selection["end"] = (event.x, event.y)
            canvas_overlay.coords(rect, event.x, event.y, event.x, event.y)

        def on_drag(event):
            if selection["start"] is None:
                return
            selection["end"] = (event.x, event.y)
            x1, y1 = selection["start"]
            x2, y2 = selection["end"]
            canvas_overlay.coords(rect, x1, y1, x2, y2)
            width = abs(x2 - x1)
            height = abs(y2 - y1)
            canvas_overlay.itemconfigure(size_hint, text=f"Size: {width} x {height}")

        def finish_selection(_event=None):
            if selection["start"] is None or selection["end"] is None:
                selection["result"] = None
            else:
                x1, y1 = selection["start"]
                x2, y2 = selection["end"]
                left, right = sorted((max(0, x1), min(base_image.width, x2)))
                top, bottom = sorted((max(0, y1), min(base_image.height, y2)))
                if right - left < 4 or bottom - top < 4:
                    selection["result"] = None
                else:
                    selection["result"] = (left, top, right, bottom)
            overlay.destroy()

        def cancel_selection(_event=None):
            selection["result"] = None
            overlay.destroy()

        canvas_overlay.bind("<ButtonPress-1>", on_press)
        canvas_overlay.bind("<B1-Motion>", on_drag)
        canvas_overlay.bind("<ButtonRelease-1>", finish_selection)
        overlay.bind("<Escape>", cancel_selection)
        overlay.grab_set()
        overlay.focus_force()
        overlay.wait_window()
        return selection["result"]

    def set_status(self, message):
        self.status_var.set(message)

    def close_app(self):
        try:
            self.note_manager.append_pending_note_to_document()
            clear_screenshot_dir()
        except Exception as exc:
            messagebox.showerror("Close failed", f"Could not finish saving/cleanup before closing.\n\n{exc}")
            return
        self.close_toast()
        self.preview_manager.close_window()
        if self.note_manager.window is not None and self.note_manager.window.winfo_exists():
            self.note_manager.window.destroy()
        self.hotkey_manager.stop()
        self.drawing_manager.hide(clear=True)
        self.root.destroy()

    def run(self):
        self.root.mainloop()


if __name__ == "__main__":
    App().run()
