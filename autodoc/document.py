import os
from datetime import datetime
from tkinter import filedialog, messagebox

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph

from .config import TEST_METADATA_FIELDS, TESTCASE_STATUS_OPTIONS
from .storage import clear_screenshot_dir, ensure_screenshot_backup, ensure_screenshot_dir, set_preview_metadata, set_preview_note
from .utils import clean_filename

class DocumentManager:
    def __init__(self, app):
        self.app = app
        self.document = None
        self.path = None
        self.tracked_images = {}
        self.tracked_notes = {}
        self.metadata_elements = []

    def has_document(self):
        return self.document is not None and self.path is not None

    def save(self):
        if not self.has_document():
            return
        try:
            self.document.save(self.path)
        except Exception as exc:
            messagebox.showerror("Save failed", f"Could not save the document.\n\n{exc}")

    def select_document(self):
        file_path = filedialog.askopenfilename(filetypes=[("Word Files", "*.docx")])
        if not file_path:
            return

        try:
            self.document = Document(file_path)
            self.path = file_path
            self.tracked_images.clear()
            self.tracked_notes.clear()
            self.metadata_elements.clear()
            imported_count = self.import_document_preview_items()
            if self.app.preview_manager.is_visible():
                self.app.preview_manager.update_canvas()
            if imported_count:
                self.app.preview_manager.refresh_paths()
                self.app.set_status(f"Selected: {os.path.basename(file_path)}. Loaded {imported_count} preview image(s).")
            else:
                self.app.set_status(f"Selected: {os.path.basename(file_path)}")
        except Exception as exc:
            messagebox.showerror("Open failed", f"Could not open document.\n\n{exc}")

    def create_new_document_file(self, folder, title):
        filename = os.path.join(folder, f"{clean_filename(title)}.docx")
        try:
            self.document = Document()
            self.document.add_heading(title, 0)
            self.document.save(filename)
            self.path = os.path.abspath(filename)
            self.tracked_images.clear()
            self.tracked_notes.clear()
            self.metadata_elements.clear()
            clear_screenshot_dir()
            self.app.preview_manager.refresh_paths()
            if self.app.preview_manager.is_visible():
                self.app.preview_manager.update_canvas()
            self.app.set_status(f"Created: {os.path.basename(self.path)}")
            return True
        except Exception as exc:
            messagebox.showerror("Create failed", f"Could not create document.\n\n{exc}")
            return False

    def _normalize_path(self, image_path):
        return os.path.abspath(image_path) if image_path else None

    def _page_content_width(self):
        section = self.document.sections[-1]
        return section.page_width - section.left_margin - section.right_margin

    def append_image(self, image_path, width=None, note_element=None):
        if not self.has_document() or not image_path:
            return

        absolute_path = self._normalize_path(image_path)
        if absolute_path in self.tracked_images:
            self.remove_tracked_image(absolute_path)

        if width is None:
            width = self._page_content_width()
        picture = self.document.add_picture(image_path, width=width)
        paragraph = self.document.paragraphs[-1]
        self.tracked_images[absolute_path] = paragraph._element
        if note_element is not None:
            self.tracked_notes[absolute_path] = note_element

    def remove_tracked_image(self, image_path, remove_note=True):
        absolute_path = self._normalize_path(image_path)
        if absolute_path is None:
            return False

        element = self.tracked_images.get(absolute_path)
        if element is None:
            return False

        parent = element.getparent()
        if parent is None:
            self.tracked_images.pop(absolute_path, None)
            return False

        parent.remove(element)
        self.tracked_images.pop(absolute_path, None)
        if remove_note:
            note_element = self.tracked_notes.pop(absolute_path, None)
            if note_element is not None:
                note_parent = note_element.getparent()
                if note_parent is not None:
                    note_parent.remove(note_element)
        return True

    def add_note_runs(self, runs):
        if not self.has_document() or not runs:
            return None

        paragraph = self.document.add_paragraph()
        for text, style in runs:
            if not text:
                continue
            run = paragraph.add_run(text)
            if style == "bold":
                run.bold = True
            elif style == "italic":
                run.italic = True
            elif style == "bold_italic":
                run.bold = True
                run.italic = True
        return paragraph._element

    def set_test_metadata(self, metadata):
        if not self.has_document():
            return False

        cleaned = {key: str(metadata.get(key, "")).strip() for key, _label in TEST_METADATA_FIELDS}
        self._remove_test_metadata_block()
        if not any(cleaned.values()):
            self.save()
            return True

        anchor = self._metadata_anchor()
        inserted_elements = []
        for key, label in reversed(TEST_METADATA_FIELDS):
            value = cleaned.get(key, "")
            if not value:
                continue
            paragraph = self.document.add_paragraph()
            label_run = paragraph.add_run(f"{label}:")
            label_run.bold = True
            if value:
                paragraph.add_run(f"\n{value}")
            paragraph.paragraph_format.space_after = Pt(6)
            element = paragraph._element
            if anchor is not None:
                anchor.addnext(element)
            inserted_elements.insert(0, element)

        self.metadata_elements = inserted_elements
        self.save()
        return True

    def get_test_metadata(self):
        metadata = {key: "" for key, _label in TEST_METADATA_FIELDS}
        if not self.has_document():
            return metadata

        for paragraph in self.document.paragraphs:
            text = paragraph.text.strip()
            for key, label in TEST_METADATA_FIELDS:
                prefix = f"{label}:"
                if text.startswith(prefix):
                    metadata[key] = text[len(prefix):].strip()
                    if paragraph._element not in self.metadata_elements:
                        self.metadata_elements.append(paragraph._element)
        return metadata

    def _metadata_anchor(self):
        for paragraph in self.document.paragraphs:
            if paragraph.text.strip():
                return paragraph._element
        return None

    def _is_test_metadata_paragraph(self, paragraph):
        text = paragraph.text.strip()
        return any(text.startswith(f"{label}:") for _key, label in TEST_METADATA_FIELDS)

    def _remove_test_metadata_block(self):
        elements = list(self.metadata_elements)
        for paragraph in self.document.paragraphs:
            if self._is_test_metadata_paragraph(paragraph) and paragraph._element not in elements:
                elements.append(paragraph._element)

        for element in elements:
            parent = element.getparent()
            if parent is not None:
                parent.remove(element)
        self.metadata_elements.clear()

    def append_testcase_status(self, status_key):
        if not self.has_document():
            return False

        status = TESTCASE_STATUS_OPTIONS.get(status_key)
        if status is None:
            return False

        spacer = self.document.add_paragraph()
        spacer.paragraph_format.space_before = Pt(8)
        spacer.paragraph_format.space_after = Pt(0)

        divider = self.document.add_paragraph()
        divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
        divider.paragraph_format.space_before = Pt(0)
        divider.paragraph_format.space_after = Pt(4)
        divider_run = divider.add_run("=" * 34)
        divider_run.bold = True
        divider_run.font.color.rgb = status["text_color"]

        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(8)

        label_run = paragraph.add_run(status["document_text"])
        label_run.bold = True
        label_run.font.size = Pt(16)
        label_run.font.color.rgb = status["text_color"]

        footer = self.document.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.paragraph_format.space_after = Pt(8)
        footer_run = footer.add_run(f"Closed on {datetime.now().strftime('%d-%m-%Y %H:%M')}")
        footer_run.italic = True
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(95, 105, 122)

        self.save()
        return True

    def set_image_note(self, image_path, note_text):
        if not self.has_document() or not image_path:
            return

        absolute_path = self._normalize_path(image_path)
        note_element = self.tracked_notes.get(absolute_path)
        if note_element is None and note_text:
            image_element = self.tracked_images.get(absolute_path)
            if image_element is not None:
                note_element = OxmlElement("w:p")
                image_element.addprevious(note_element)
                self.tracked_notes[absolute_path] = note_element

        if note_element is None:
            return

        paragraph = Paragraph(note_element, self.document)
        paragraph.clear()
        if note_text:
            paragraph.add_run(note_text)
        self.save()

    def import_document_preview_items(self):
        if not self.has_document():
            return 0

        clear_screenshot_dir()
        set_preview_metadata(self.get_test_metadata())
        imported_count = 0
        pending_note_parts = []

        for paragraph in self.document.paragraphs:
            image_rids = paragraph._element.xpath(".//a:blip/@r:embed")
            if not image_rids:
                if self._is_test_metadata_paragraph(paragraph):
                    continue
                text = paragraph.text.strip()
                if text:
                    pending_note_parts.append(text)
                continue

            note_text = "\n".join(pending_note_parts).strip()
            pending_note_parts = []

            for rid in image_rids:
                related_part = self.document.part.related_parts.get(rid)
                if related_part is None:
                    continue
                extension = os.path.splitext(related_part.partname)[1] or ".png"
                imported_count += 1
                image_name = f"doc_image_{imported_count:03d}{extension}"
                image_path = os.path.join(ensure_screenshot_dir(), image_name)
                try:
                    with open(image_path, "wb") as image_file:
                        image_file.write(related_part.blob)
                    ensure_screenshot_backup(image_path)
                except OSError:
                    continue

                absolute_path = self._normalize_path(image_path)
                self.tracked_images[absolute_path] = paragraph._element
                if note_text:
                    self.tracked_notes[absolute_path] = self._find_note_element_before(paragraph._element)
                    set_preview_note(image_path, note_text)

        return imported_count

    def _find_note_element_before(self, image_element):
        previous = image_element.getprevious()
        while previous is not None:
            paragraph = Paragraph(previous, self.document)
            if paragraph.text.strip():
                return previous
            previous = previous.getprevious()
        return None


